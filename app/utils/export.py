"""
Export utilities for generating PDF and Word documents from markdown reports.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

try:
    from weasyprint import HTML, CSS
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
    print("Warning: WeasyPrint not available. PDF export will use alternative method.")

import markdown
import os
import re


def generate_pdf(analysis_id: str) -> str:
    """
    Generate a styled PDF from the markdown report.

    Args:
        analysis_id: Analysis ID

    Returns:
        Path to generated PDF file

    Raises:
        FileNotFoundError: If report markdown doesn't exist
        RuntimeError: If WeasyPrint is not available
    """
    if not WEASYPRINT_AVAILABLE:
        raise RuntimeError(
            "PDF generation not available. WeasyPrint requires additional system libraries. "
            "Please use Word export or Markdown copy instead. "
            "See: https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#installation"
        )

    report_path = f"analyses/{analysis_id}/report.md"
    pdf_path = f"analyses/{analysis_id}/report.pdf"

    # Read markdown content
    with open(report_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    # Convert markdown to HTML with extensions
    html_content = markdown.markdown(
        md_content,
        extensions=['tables', 'fenced_code', 'nl2br']
    )

    # Create styled HTML
    styled_html = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <style>
            @page {{
                margin: 2cm;
                size: A4;
            }}

            body {{
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                font-size: 11pt;
                line-height: 1.6;
                color: #333;
            }}

            h1 {{
                color: #1a1a1a;
                border-bottom: 3px solid #4A90E2;
                padding-bottom: 10px;
                margin-top: 20px;
                font-size: 24pt;
            }}

            h2 {{
                color: #2c3e50;
                border-bottom: 2px solid #ddd;
                padding-bottom: 5px;
                margin-top: 20px;
                font-size: 18pt;
            }}

            h3 {{
                color: #34495e;
                margin-top: 15px;
                font-size: 14pt;
            }}

            code {{
                background-color: #f4f4f4;
                padding: 2px 6px;
                border-radius: 3px;
                font-family: 'Courier New', monospace;
                font-size: 10pt;
            }}

            pre {{
                background-color: #f8f8f8;
                border: 1px solid #ddd;
                border-left: 4px solid #4A90E2;
                padding: 15px;
                overflow-x: auto;
                border-radius: 4px;
            }}

            pre code {{
                background: none;
                padding: 0;
            }}

            table {{
                border-collapse: collapse;
                width: 100%;
                margin: 20px 0;
            }}

            th, td {{
                border: 1px solid #ddd;
                padding: 10px;
                text-align: left;
            }}

            th {{
                background-color: #4A90E2;
                color: white;
                font-weight: bold;
            }}

            tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}

            strong {{
                color: #2c3e50;
            }}

            em {{
                color: #555;
            }}

            ul, ol {{
                margin: 10px 0;
                padding-left: 30px;
            }}

            li {{
                margin: 5px 0;
            }}

            hr {{
                border: none;
                border-top: 1px solid #ddd;
                margin: 20px 0;
            }}

            blockquote {{
                border-left: 4px solid #4A90E2;
                margin: 15px 0;
                padding-left: 15px;
                color: #666;
                font-style: italic;
            }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """

    # Generate PDF using weasyprint
    HTML(string=styled_html).write_pdf(pdf_path)

    return pdf_path


def generate_docx(analysis_id: str) -> str:
    """
    Generate a Word document from the markdown report.

    Args:
        analysis_id: Analysis ID

    Returns:
        Path to generated Word document

    Raises:
        FileNotFoundError: If report markdown doesn't exist
    """
    report_path = f"analyses/{analysis_id}/report.md"
    docx_path = f"analyses/{analysis_id}/report.docx"

    # Read markdown content
    with open(report_path, "r", encoding="utf-8") as f:
        md_content = f.read()

    # Create Document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Parse markdown line by line (basic parser)
    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Heading 1
        if line.startswith('# '):
            heading = doc.add_heading(line[2:].strip(), level=1)
            heading.style.font.color.rgb = RGBColor(26, 26, 26)

        # Heading 2
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:].strip(), level=2)
            heading.style.font.color.rgb = RGBColor(44, 62, 80)

        # Heading 3
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:].strip(), level=3)
            heading.style.font.color.rgb = RGBColor(52, 73, 94)

        # Horizontal rule
        elif line.startswith('---') or line.startswith('***'):
            doc.add_paragraph('_' * 60)

        # Unordered list
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'`(.*?)`', r'\1', text)  # Code
            doc.add_paragraph(text, style='List Bullet')

        # Ordered list
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            doc.add_paragraph(text, style='List Number')

        # Code block
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style.font.name = 'Courier New'
            p.style.font.size = Pt(9)

        # Regular paragraph
        else:
            # Remove markdown formatting while adding to doc
            text = line

            # Simple bold and italic handling
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)

            # Add paragraph
            if text.strip():
                doc.add_paragraph(text.strip())

        i += 1

    # Save document
    doc.save(docx_path)

    return docx_path
